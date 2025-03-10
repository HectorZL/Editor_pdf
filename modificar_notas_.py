import pdfplumber
import pandas as pd
from docx import Document
from pdf2docx import Converter
import os
import win32com.client  # Import win32com for Word to PDF conversion
import lxml.etree as ET  # Cambié esto para usar lxml.etree

from docx.shared import Pt  # Importar para definir el tamaño de la fuente

# Diccionario de notas posibles
notas_posibles = {
    "NE": "NE",  # Nota para borrar
    "A+": "A+", 
    "A-": "A-", 
    "B+": "B+", 
    "B-": "B-", 
    "C+": "C+", 
    "C-": "C-", 
    "D+": "D+", 
    "D-": "D-", 
    "E+": "E+", 
    "E-": "E-"
}

def pdf_a_word(pdf_path, word_path):
    # Convertir el PDF a Word manteniendo la estructura de las tablas
    cv = Converter(pdf_path)
    cv.convert(word_path, start=0, end=None)
    print(f"📄 Documento Word creado en {word_path}")

    # Intentar agregar bordes a las tablas si no se preservaron
    doc = Document(word_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Añadir borde a las celdas si no existe
                tcPr = cell._element.get_or_add_tcPr()
                if not tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders"):
                    borders = ET.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders")
                    tcPr.append(borders)
    doc.save(word_path)
    print(f"✅ Bordes de tabla añadidos en {word_path}")

def modificar_notas_word(word_path, notas_trimestres):
    """
    Modificar las notas en el archivo Word según las notas ingresadas por trimestre.
    - notas_trimestres: diccionario con las notas a asignar para cada trimestre (1, 2, 3).
    """
    # Abrir el archivo Word
    doc = Document(word_path)

    # Recorrer las tablas en el documento Word
    for table in doc.tables:
        for row in table.rows:
            # Recorrer las celdas en cada fila
            for i, cell in enumerate(row.cells):
                # Si se encuentra alguna nota "NE", reemplazarla
                for nota, nueva_nota in notas_posibles.items():
                    if nota in cell.text:
                        # Asignar las notas ingresadas por el usuario según el trimestre
                        if i == 2 and 1 in notas_trimestres:  # Columna 3 (primer trimestre)
                            nueva_nota = notas_trimestres.get(1, nueva_nota)
                            cell.text = cell.text.replace(nota, nueva_nota)
                        elif i == 3 and 2 in notas_trimestres:  # Columna 4 (segundo trimestre)
                            nueva_nota = notas_trimestres.get(2, nueva_nota)
                            cell.text = cell.text.replace(nota, nueva_nota)
                        elif i == 4 and 3 in notas_trimestres:  # Columna 5 (tercer trimestre)
                            nueva_nota = notas_trimestres.get(3, nueva_nota)
                            cell.text = cell.text.replace(nota, nueva_nota)

                        # Cambiar el tamaño de la fuente para hacerlo 1 tipo menos
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)  # Cambiar el tamaño de la fuente (ajustar según sea necesario)

                        # Centrar el texto en la celda
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = 1  # 1 es para centrar el texto

    # Guardar el documento con las notas modificadas
    doc.save(word_path)
    print(f"✅ Notas modificadas y guardadas en {word_path}")

def word_a_pdf(word_path, pdf_output):
    # Verificar si el archivo existe
    if not os.path.exists(word_path):
        print(f"❌ El archivo {word_path} no se encuentra en la ubicación especificada.")
        return
    
    # Mostrar el camino completo del archivo para depuración
    print(f"🔍 Intentando abrir el archivo Word: {word_path}")

    try:
        # Convertir el documento Word a PDF (utilizando win32com)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # Para que no se muestre la aplicación de Word
        word.DisplayAlerts = False  # Evitar alertas

        # Espera a que Word termine de abrir
        doc = word.Documents.Open(word_path)

        # Guardar como PDF
        doc.SaveAs(pdf_output, FileFormat=17)  # FileFormat=17 corresponde a PDF
        doc.Close()
        word.Quit()

        print(f"📂 PDF generado correctamente en {pdf_output}")
    
    except Exception as e:
        print(f"❌ Error al convertir de Word a PDF: {e}")

def procesar_lote_notas(base_folder, output_folder, notas_trimestres):
    # Verificar si las carpetas existen
    if not os.path.exists(base_folder):
        print(f"❌ La carpeta de entrada {base_folder} no existe.")
        return
    
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)  # Crear la carpeta de salida si no existe

    # Obtener todos los archivos PDF en la carpeta de entrada
    for filename in os.listdir(base_folder):
        if filename.endswith(".pdf"):  # Filtrar solo archivos PDF
            pdf_path = os.path.join(base_folder, filename)
            word_path = os.path.join(base_folder, f"{os.path.splitext(filename)[0]}.docx")
            pdf_output = os.path.join(output_folder, f"{os.path.splitext(filename)[0]}_corregido.pdf")

            # Procesar el archivo
            print(f"\nProcesando el archivo: {filename}")
            pdf_a_word(pdf_path, word_path)  # Convertir PDF a Word
            modificar_notas_word(word_path, notas_trimestres)  # Modificar las notas
            word_a_pdf(word_path, pdf_output)  # Convertir el Word a PDF

            # Borrar el archivo temporal de Word
            if os.path.exists(word_path):
                os.remove(word_path)
                print(f"❌ Archivo temporal Word borrado: {word_path}")

# Función para pedir al usuario las notas para cada trimestre
def pedir_notas():
    print("Ingrese las notas para cada trimestre (deje vacío para no modificar, 'NE' para borrar):")
    notas = {}
    for trimestre in [1, 2, 3]:
        nota = input(f"Nota para el {trimestre}º trimestre: ")
        if nota:  # Si no está vacío, se asigna la nota ingresada
            if nota in notas_posibles.values():  # Validar que la nota ingresada sea válida
                notas[trimestre] = nota
            else:
                print("❌ Nota no válida. Las opciones son: A+, A-, B+, B-, C+, C-, D+, D-, E+, E-, o 'NE' para borrar.")
                return {}
    return notas

# Rutas de carpetas
base_folder = r"C:\Users\jesuc\Documents\EditorPdf\notas_base"  # Carpeta de entrada
output_folder = r"C:\Users\jesuc\Documents\EditorPdf\notas_correjido"  # Carpeta de salida

# Pedir las notas al usuario
notas_trimestres = pedir_notas()

# Si las notas son válidas, procesar los archivos
if notas_trimestres:
    procesar_lote_notas(base_folder, output_folder, notas_trimestres)
