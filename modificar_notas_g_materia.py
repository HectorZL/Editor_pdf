import pdfplumber
import pandas as pd
from docx import Document
from pdf2docx import Converter
import os
import win32com.client  # Import win32com for Word to PDF conversion
import lxml.etree as ET  # Cambié esto para usar lxml.etree

from docx.shared import Pt  # Importar para definir el tamaño de la fuente

def pdf_a_word(pdf_path, word_path):
    # Convertir el PDF a Word manteniendo la estructura de las tablas
    cv = Converter(pdf_path)
    cv.convert(word_path, start=0, end=None)
    print(f"📄 Documento Word creado en {word_path}")

    # Intentar agregar bordes a las tablas si no se preservaron
    doc = Document(word_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Añadir borde a las celdas si no existe
                tcPr = cell._element.get_or_add_tcPr()
                if not tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders"):
                    borders = ET.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders")
                    tcPr.append(borders)
    doc.save(word_path)
    print(f"✅ Bordes de tabla añadidos en {word_path}")

def modificar_notas_word(word_path):
    # Abrir el archivo Word
    doc = Document(word_path)

    # Recorrer las tablas en el documento Word
    for table in doc.tables:
        for row in table.rows:
            # Verificar si la materia "COMPRENSIÓN Y EXPRESIÓN DEL LENGUAJE" está en la primera columna
            if "COMPRENSIÓN Y EXPRESIÓN DEL LENGUAJE" in row.cells[0].text:
                # Cambiar todas las celdas correspondientes a los trimestres (1er, 2do y 3ro) por "C+"
                for i in range(3, 5):  # Las columnas 3, 4 y 5 corresponden a los trimestres
                    row.cells[i].text = "A+"  # Cambiar el texto de cada celda a "C+"
                    
                    # Cambiar el tamaño de la fuente para hacerlo 1 tipo menos
                    for paragraph in row.cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)  # Cambiar el tamaño de la fuente (ajustar según sea necesario)

                    # Centrar el texto en la celda
                    for paragraph in row.cells[i].paragraphs:
                        paragraph.alignment = 1  # 1 es para centrar el texto

    # Guardar el documento con las notas modificadas
    doc.save(word_path)
    print(f"✅ Notas modificadas y guardadas en {word_path}")



def word_a_pdf(word_path, pdf_output):
    # Verificar si el archivo existe
    if not os.path.exists(word_path):
        print(f"❌ El archivo {word_path} no se encuentra en la ubicación especificada.")
        return
    
    # Mostrar el camino completo del archivo para depuración
    print(f"🔍 Intentando abrir el archivo Word: {word_path}")

    try:
        # Convertir el documento Word a PDF (utilizando win32com)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # Para que no se muestre la aplicación de Word
        word.DisplayAlerts = False  # Evitar alertas

        # Espera a que Word termine de abrir
        doc = word.Documents.Open(word_path)

        # Guardar como PDF
        doc.SaveAs(pdf_output, FileFormat=17)  # FileFormat=17 corresponde a PDF
        doc.Close()
        word.Quit()

        print(f"📂 PDF generado correctamente en {pdf_output}")
    
    except Exception as e:
        print(f"❌ Error al convertir de Word a PDF: {e}")

def procesar_lote_notas(base_folder, output_folder):
    # Verificar si las carpetas existen
    if not os.path.exists(base_folder):
        print(f"❌ La carpeta de entrada {base_folder} no existe.")
        return
    
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)  # Crear la carpeta de salida si no existe

    # Obtener todos los archivos PDF en la carpeta de entrada
    for filename in os.listdir(base_folder):
        if filename.endswith(".pdf"):  # Filtrar solo archivos PDF
            pdf_path = os.path.join(base_folder, filename)
            word_path = os.path.join(base_folder, f"{os.path.splitext(filename)[0]}.docx")
            pdf_output = os.path.join(output_folder, f"{os.path.splitext(filename)[0]}_corregido.pdf")

            # Procesar el archivo
            print(f"\nProcesando el archivo: {filename}")
            pdf_a_word(pdf_path, word_path)  # Convertir PDF a Word
            modificar_notas_word(word_path)  # Modificar las notas
            word_a_pdf(word_path, pdf_output)  # Convertir el Word a PDF

            # Borrar el archivo temporal de Word
            if os.path.exists(word_path):
                os.remove(word_path)
                print(f"❌ Archivo temporal Word borrado: {word_path}")

# Rutas de carpetas
base_folder = r"C:\Users\jesuc\Documents\EditorPdf\notas_base_1"  # Carpeta de entrada
output_folder = r"C:\Users\jesuc\Documents\EditorPdf\notas_correjido_1"  # Carpeta de salida

# Procesar todos los archivos en la carpeta de entrada
procesar_lote_notas(base_folder, output_folder)
